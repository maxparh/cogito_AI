from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
import os
import random

from models import Ticket, SubmitAnswerRequest
from storage import tickets_db, current_ticket_id
from ai import extract_tickets_from_text, call_ollama
from docx import Document
from io import BytesIO

from typing import List
from pydantic import BaseModel

# Pydantic модели для новых данных
class Lesson(BaseModel):
    day: str
    subject: str
    time: str
    room: str

class Grade(BaseModel):
    name: str
    attendance: int
    score: int
    maxScore: int

class TutorSession(BaseModel):
    subject: str
    datetime: str
    tutor: str
    price: str

# Создаём приложение
app = FastAPI()

# Главная страница — отдаём HTML
@app.get("/", response_class=HTMLResponse)
async def home():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

# --- Новые эндпоинты для дашборда ---

@app.get("/api/schedule/", response_model=List[Lesson])
async def get_schedule():
    data = [
        {"day": "Понедельник", "subject": "Алгоритмизация и программирование", "time": "11:00", "room": "453"},
        {"day": "Понедельник", "subject": "ИСИТ", "time": "12:40", "room": "487"},
        {"day": "Вторник", "subject": "Алгоритмизация и программирование", "time": "11:00", "room": "453"},
        {"day": "Вторник", "subject": "ИСИТ", "time": "12:40", "room": "487"},
        {"day": "Среда", "subject": "Алгоритмизация и программирование", "time": "11:00", "room": "453"},
        {"day": "Среда", "subject": "ИСИТ", "time": "12:40", "room": "487"},
        {"day": "Четверг", "subject": "Алгоритмизация и программирование", "time": "11:00", "room": "453"},
        {"day": "Четверг", "subject": "ИСИТ", "time": "12:40", "room": "487"},
        {"day": "Четверг", "subject": "Веб-дизайн", "time": "12:40", "room": "487"},
        {"day": "Пятница", "subject": "Алгоритмизация и программирование", "time": "11:00", "room": "453"},
        {"day": "Пятница", "subject": "ИСИТ", "time": "12:40", "room": "487"},
        {"day": "Пятница", "subject": "Веб-дизайн", "time": "12:40", "room": "487"}
    ]
    return data

@app.get("/api/grades/", response_model=List[Grade])
async def get_grades():
    data = [
        {"name": "Алгоритмизация и программирование", "attendance": 70, "score": 68, "maxScore": 100},
        {"name": "ИСИТ", "attendance": 70, "score": 68, "maxScore": 100},
        {"name": "ВЕБ-дизайн", "attendance": 70, "score": 68, "maxScore": 100}
    ]
    return data

@app.get("/api/tutor-sessions/", response_model=List[TutorSession])
async def get_tutor_sessions():
    data = [
        {"subject": "Алгоритмизация и программирование", "datetime": "23.11 19:00", "tutor": "Марина Иванова", "price": "800"}
    ]
    return data

# --- Конец новых эндпоинтов ---

# Загрузка файла
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()

    # Проверяем расширение файла
    if file.filename.endswith(".docx"):
        # Читаем .docx
        doc = Document(BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        # Для .txt или других — декодируем как текст
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("cp1251")
            except UnicodeDecodeError:
                text = content.decode("utf-8", errors="replace")

    global tickets_db
    tickets_db = extract_tickets_from_text(text)

    return {
        "message": f"Файл обработан. Сгенерировано {len(tickets_db)} билетов.",
        "tickets": tickets_db
    }

# Получить случайный билет
@app.get("/get_ticket", response_model=Ticket)
async def get_random_ticket():
    if not tickets_db:
        return {"error": "Нет загруженных билетов. Сначала загрузите файл."}

    ticket = random.choice(tickets_db)
    return ticket

# Отправить ответ на вопрос
@app.post("/submit_answer")
async def submit_answer(request: SubmitAnswerRequest):
    # Найдём билет и вопрос
    ticket = next((t for t in tickets_db if t["ticket_id"] == request.ticket_id), None)
    if not ticket:
        return {"error": "Билет не найден"}

    question = next((q for q in ticket["questions"] if q["question_id"] == request.question_id), None)
    if not question:
        return {"error": "Вопрос не найден"}

    # Пока эталонный ответ — заглушка
    correct_answer = "Представь, что ты знаешь правильный ответ."
    prompt = f"""
    Оцени, насколько точно пользователь ответил на вопрос.
    Ответь кратко: 'точно', 'частично', 'нет' или 'не по теме'.
    Также добавь одну строку объяснения.

    Вопрос: {question['question_text']}
    Задание: {question.get('task', 'нет')}
    Эталонный ответ: {correct_answer}
    Ответ пользователя: {request.user_answer}

    Формат вывода:
    Оценка: [точно/частично/нет/не по теме]
    Объяснение: [текст]
    """

    feedback = call_ollama(prompt)

    return {
        "ticket_id": request.ticket_id,
        "question_id": request.question_id,
        "user_answer": request.user_answer,
        "ai_feedback": feedback.strip()
    }

@app.post("/extract_text")
async def extract_text(file: UploadFile = File(...)):
    """Извлекает чистый текст из .docx или .txt файла"""
    content = await file.read()

    if file.filename.endswith(".docx"):
        doc = Document(BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("cp1251")
            except UnicodeDecodeError:
                text = content.decode("utf-8", errors="replace")
    
    return {"original_text": text}


@app.post("/research_advice")
async def get_research_advice(request: dict):
    text = request.get("original_text", "")  # ВАЖНО: original_text, не content!
    if not text:
        return {"error": "Текст для анализа не найден."}

    prompt = f"""
Ты — эксперт по академическому письму. Твоя задача — дать 3–5 практических советов по улучшению текста.

Правила:
- Отвечай ТОЛЬКО советами, каждый с новой строки.
- НЕ пиши вводные фразы: "Вот советы", "Рекомендую", "Советую", "Я считаю".
- НЕ используй маркированные списки (типа * или 1.), НЕ используй цифры, НЕ используй тире.
- НЕ оборачивай ответ в JSON, НЕ используй кавычки, НЕ используй скобки.
- Ответ должен быть чистым текстом — как будто ты говоришь студенту в лицо.

Текст для анализа:
{text}
"""

    advice = call_ollama(prompt)
    return {"advice": advice.strip()}


@app.post("/research_antigpt")
async def check_antigpt(request: dict):
    text = request.get("original_text", "")  # ВАЖНО: original_text, не content!
    if not text:
        return {"error": "Текст для проверки не найден."}

    prompt = f"""
Ты — эксперт по обнаружению генерации текста ИИ. Проанализируй текст и дай ответ строго в формате:

Вероятность генерации ИИ: X%
Анализ: [1–2 предложения о том, что выглядит подозрительно, и почему]

Правила:
- НЕ пиши ничего, кроме этого формата.
- НЕ добавляй пояснений, НЕ используй слова "возможно", "кажется", "вероятно".
- Если текст выглядит человеческим — напиши: "Вероятность генерации ИИ: 5%" и "Анализ: Текст содержит личные формулировки и нестандартные конструкции, характерные для человека."
- НЕ используй JSON, НЕ используй списки, НЕ используй маркеры.
- Ответ должен быть ровно 2 строки — без лишних символов.

Текст:
{text}
"""

    result = call_ollama(prompt)
    return {"result": result.strip()}
# Подключаем статику — должно быть в конце
app.mount("/static", StaticFiles(directory="static"), name="static")